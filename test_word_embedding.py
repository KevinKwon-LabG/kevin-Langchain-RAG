#!/usr/bin/env python3
"""
워드 임베딩 기능 테스트 스크립트
"""

import os
import sys
import logging
from pathlib import Path

# 프로젝트 루트를 Python 경로에 추가
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from src.services.word_embedding_service import WordEmbeddingService

# 로깅 설정
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


def create_test_word_document():
    """테스트용 워드 문서 생성"""
    try:
        from docx import Document
        
        # 테스트 문서 생성
        doc = Document()
        
        # 제목 추가
        doc.add_heading('테스트 문서', 0)
        
        # 단락 추가
        doc.add_paragraph('이것은 워드 임베딩 테스트를 위한 문서입니다.')
        doc.add_paragraph('한국어 자연어 처리를 위한 형태소 분석이 적용됩니다.')
        
        # 소제목 추가
        doc.add_heading('주요 기능', level=1)
        
        # 리스트 추가
        features = [
            '텍스트 추출: python-docx를 사용한 워드 문서 텍스트 추출',
            '전처리: 텍스트 정제 및 형태소 분석',
            '문서 분할: 100~500 토큰 단위로 분할',
            '임베딩 생성: KoBERT를 사용한 벡터화',
            '벡터 검색: 유사한 문서 청크 검색'
        ]
        
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')
        
        # 테이블 추가
        doc.add_heading('처리 단계', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # 헤더 설정
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '단계'
        hdr_cells[1].text = '설명'
        hdr_cells[2].text = '도구'
        
        # 데이터 추가
        steps = [
            ('1', '텍스트 추출', 'python-docx'),
            ('2', '전처리', 'KoNLPy (Okt, Mecab, Komoran)'),
            ('3', '문서 분할', 'tiktoken'),
            ('4', '임베딩 생성', 'Sentence Transformers'),
            ('5', '벡터 저장', 'ChromaDB')
        ]
        
        for step_num, description, tool in steps:
            row_cells = table.add_row().cells
            row_cells[0].text = step_num
            row_cells[1].text = description
            row_cells[2].text = tool
        
        # 테스트 디렉토리 생성
        test_dir = Path('./data/test')
        test_dir.mkdir(parents=True, exist_ok=True)
        
        # 문서 저장
        test_file_path = test_dir / 'test_document.docx'
        doc.save(str(test_file_path))
        
        print(f"    테스트 워드 문서 생성 완료: {test_file_path}")
        return str(test_file_path)
        
    except Exception as e:
        print(f"    테스트 문서 생성 실패: {e}")
        return None


def test_word_embedding_service():
    """워드 임베딩 서비스 테스트"""
    try:
        print("  워드 임베딩 서비스 테스트 시작")
        
        # 1. 서비스 초기화
        print("  1. 서비스 초기화")
        service = WordEmbeddingService(
            chunk_size=200,
            chunk_overlap=30
        )
        print(f"    ✅ 서비스 초기화 완료 (모델: {service.embedding_model_name})")
        
        # 2. 테스트 문서 생성
        print("  2. 테스트 문서 생성")
        test_file_path = create_test_word_document()
        if not test_file_path:
            print("    ❌ 테스트 문서 생성 실패")
            return False
        print(f"    ✅ 테스트 문서 생성 완료: {test_file_path}")
        
        # 3. 문서 처리
        print("  3. 문서 처리")
        result = service.process_word_document(
            test_file_path,
            metadata={'test': True, 'category': 'test'}
        )
        
        if result['processing_success']:
            print(f"    ✅ 문서 처리 성공: {result['total_chunks']}개 청크, {result['total_tokens']}개 토큰")
        else:
            print(f"    ❌ 문서 처리 실패: {result.get('error', 'Unknown error')}")
            return False
        
        # 4. 검색 테스트
        print("  4. 검색 테스트")
        search_queries = [
            "형태소 분석",
            "임베딩 생성",
            "문서 분할",
            "벡터 검색"
        ]
        
        for query in search_queries:
            print(f"    검색 쿼리: '{query}'")
            results = service.search_similar_chunks(query, n_results=3)
            
            for i, result in enumerate(results):
                print(f"      결과 {i+1}: {result['content'][:100]}...")
                print(f"        거리: {result['distance']:.4f}")
                print(f"        파일: {result['metadata'].get('file_name', 'Unknown')}")
        
        # 5. 통계 확인
        print("  5. 통계 확인")
        stats = service.get_collection_stats()
        print(f"    컬렉션 통계: {stats}")
        
        print("  ✅ 워드 임베딩 서비스 테스트 완료")
        return True
        
    except Exception as e:
        print(f"  ❌ 워드 임베딩 서비스 테스트 실패: {e}")
        return False


def test_api_endpoints():
    """API 엔드포인트 테스트"""
    try:
        import requests
        import json
        
        # 환경 설정에서 서비스 URL 가져오기
        try:
            from src.config.settings import get_settings
            settings = get_settings()
            base_url = settings.service_url
        except:
            base_url = "http://1.237.52.240:11040"
        
        logger.info("API 엔드포인트 테스트 시작")
        
        # 1. 헬스 체크
        logger.info("1. 헬스 체크")
        response = requests.get(f"{base_url}/api/word-embedding/health")
        if response.status_code == 200:
            health_data = response.json()
            logger.info(f"헬스 체크 성공: {health_data}")
        else:
            logger.error(f"헬스 체크 실패: {response.status_code}")
            return False
        
        # 2. 통계 조회
        logger.info("2. 통계 조회")
        response = requests.get(f"{base_url}/api/word-embedding/stats")
        if response.status_code == 200:
            stats_data = response.json()
            logger.info(f"통계 조회 성공: {stats_data}")
        else:
            logger.error(f"통계 조회 실패: {response.status_code}")
        
        # 3. 검색 테스트
        logger.info("3. 검색 테스트")
        search_data = {
            "query": "형태소 분석",
            "n_results": 3
        }
        
        response = requests.post(
            f"{base_url}/api/word-embedding/search",
            json=search_data
        )
        
        if response.status_code == 200:
            search_results = response.json()
            logger.info(f"검색 성공: {len(search_results)}개 결과")
            for i, result in enumerate(search_results):
                logger.info(f"  결과 {i+1}: {result['content'][:100]}...")
        else:
            logger.error(f"검색 실패: {response.status_code}")
        
        logger.info("API 엔드포인트 테스트 완료")
        return True
        
    except Exception as e:
        logger.error(f"API 엔드포인트 테스트 실패: {e}")
        return False


if __name__ == "__main__":
    print("=" * 60)
    print("워드 임베딩 테스트 시작")
    print("=" * 60)
    
    # 서비스 테스트
    print("\n1. 워드 임베딩 서비스 테스트")
    service_test_success = test_word_embedding_service()
    
    if service_test_success:
        print("✅ 서비스 테스트 성공")
        
        # API 테스트 (서버가 실행 중인 경우)
        print("\n2. API 엔드포인트 테스트")
        try:
            api_test_success = test_api_endpoints()
            if api_test_success:
                print("✅ API 테스트 성공")
            else:
                print("⚠️ API 테스트 실패 (서버가 실행되지 않았을 수 있음)")
        except Exception as e:
            print(f"⚠️ API 테스트 건너뜀: {e}")
    else:
        print("❌ 서비스 테스트 실패")
    
    print("\n" + "=" * 60)
    print("워드 임베딩 테스트 완료")
    print("=" * 60) 