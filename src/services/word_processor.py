#!/usr/bin/env python3
"""
워드 문서 텍스트 추출 유틸리티
RAG 시스템에서 워드 문서의 텍스트를 추출하는 기본 기능
"""

import logging
import re
from typing import Dict, Any
from pathlib import Path

# 워드 문서 처리 라이브러리
from docx import Document

logger = logging.getLogger(__name__)

class WordProcessor:
    """워드 문서 텍스트 추출 클래스"""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.doc']
    
    def extract_text_from_word(self, file_path: str) -> str:
        """
        워드 문서에서 텍스트 추출
        
        Args:
            file_path: 워드 파일 경로
            
        Returns:
            str: 추출된 텍스트
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"지원하지 않는 워드 파일 형식: {file_extension}")
            
            logger.info(f"워드 파일 텍스트 추출 시작: {file_path}")
            
            doc = Document(file_path)
            text_parts = []
            
            # 단락에서 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 테이블에서 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            extracted_text = "\n".join(text_parts)
            
            # 불필요한 요소 제거
            cleaned_text = self._clean_text(extracted_text)
            
            logger.info(f"텍스트 추출 완료: {len(cleaned_text)} 문자")
            
            return cleaned_text
                
        except Exception as e:
            logger.error(f"워드 파일 텍스트 추출 실패: {file_path}, 오류: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """
        텍스트에서 불필요한 요소 제거
        
        Args:
            text: 원본 텍스트
            
        Returns:
            str: 정제된 텍스트
        """
        if not text:
            return ""
        
        # 1. 기본 정리
        text = text.strip()
        
        # 2. 줄 바꿈, 탭을 공백으로 변환
        text = re.sub(r'[\n\t\r]+', ' ', text)
        
        # 3. 특수 문자 정제
        # 다양한 대시 문자를 하이픈으로 통일
        text = re.sub(r'[—–−]', '-', text)
        
        # 다양한 불릿 문자 통일
        text = re.sub(r'[•·∙]', '•', text)
        
        # 다양한 따옴표 통일
        text = re.sub(r'[""""]', '"', text)
        text = re.sub(r"[''']", "'", text)
        
        # 기타 특수 문자 제거 (한글, 영문, 숫자, 기본 문장부호 제외)
        text = re.sub(r'[^\w\s\.,!?;:()\-_가-힣•"\'"]', ' ', text)
        
        # 4. 반복된 공백 제거
        text = re.sub(r'\s+', ' ', text)
        
        # 5. 앞뒤 공백 제거
        text = text.strip()
        
        # 6. 연속된 문장부호 정리
        text = re.sub(r'[.,!?;:]{2,}', lambda m: m.group()[0], text)
        
        # 7. 공백과 문장부호 사이 정리
        text = re.sub(r'\s+([.,!?;:])', r'\1', text)
        
        # 8. 괄호 내부 공백 정리
        text = re.sub(r'\(\s+', '(', text)
        text = re.sub(r'\s+\)', ')', text)
        
        # 9. 최종 정리
        text = text.strip()
        
        return text

# 전역 인스턴스
word_processor = WordProcessor() 